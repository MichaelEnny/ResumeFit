"""Converts the ATS-safe plain-text rewrite into a downloadable, already-formatted
DOCX file - not just plain paragraphs the user then has to reformat by hand.

The rewriter's plain text has a predictable shape: a name line, a contact line,
then Section headers, and within Experience/Education a "block" per job or degree
- a title line, optionally one or two more plain lines (company, dates), then
"- " bullets. This module recognizes that shape instead of treating every plain
line identically, so the DOCX comes out with real hierarchy: a bold centered name,
a muted contact line, accent-colored section headers with a rule under them, and
bold job-title lines paired with an italic company/dates line - the layout a
person would otherwise have to build by hand.

None of this touches ATS safety: bold, italic, color, and paragraph borders are
purely cosmetic to a text-extracting parser (unlike the tables/images/text boxes
src/document_parser.py already flags as risky), so the structural guarantees the
rewrite prompt already makes still hold.
"""
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

SECTION_HEADERS = {"summary", "experience", "education", "skills"}
# Entries within these sections get the bold-title / italic-meta treatment.
# Summary is prose, Skills is just a bullet list - both stay plain.
ENTRY_SECTIONS = {"experience", "education"}

# Education entries don't reliably list the degree before the institution (or
# vice versa), unlike Experience entries where the role is always first. Rather
# than always bolding whichever line happens to come first, search for a
# degree-shaped line and bold that one instead.
DEGREE_KEYWORDS = (
    "bachelor", "master", "associate", "ph.d", "phd", "b.a.", "b.s.",
    "m.a.", "m.s.", "m.b.a", "diploma", "certificate", "doctorate",
)

ACCENT = RGBColor(0x0B, 0x6E, 0x52)
MUTED = RGBColor(0x55, 0x53, 0x4C)
BASE_FONT = "Calibri"


def _is_section_header(stripped: str) -> bool:
    return stripped.rstrip(":").lower() in SECTION_HEADERS


def _looks_like_degree(line: str) -> bool:
    lowered = line.lower()
    return any(keyword in lowered for keyword in DEGREE_KEYWORDS)


def _add_bottom_rule(paragraph, color_hex: str, size: int = 6) -> None:
    """Adds a thin bottom border to a paragraph, used under section headers."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color_hex)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def build_docx(resume_text: str) -> bytes:
    document = Document()
    base_style = document.styles["Normal"]
    base_style.font.name = BASE_FONT
    base_style.font.size = Pt(10.5)

    lines = resume_text.splitlines()
    n = len(lines)
    i = 0

    # --- Header block: first non-empty lines before the first section header
    # are the name (bold, larger, centered) then contact info (muted, centered).
    header_line_index = 0
    while i < n and lines[i].strip() and not _is_section_header(lines[i].strip()):
        stripped = lines[i].strip()
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(stripped)
        if header_line_index == 0:
            run.bold = True
            run.font.size = Pt(20)
            p.paragraph_format.space_after = Pt(2)
        else:
            run.font.size = Pt(10)
            run.font.color.rgb = MUTED
            p.paragraph_format.space_after = Pt(2)
        header_line_index += 1
        i += 1
    while i < n and not lines[i].strip():
        i += 1

    current_section = None
    entry_buffer: list[str] = []

    def flush_entry() -> None:
        if not entry_buffer:
            return
        title_index = 0
        if current_section == "education":
            for idx, line in enumerate(entry_buffer):
                if _looks_like_degree(line):
                    title_index = idx
                    break
        title_line = entry_buffer[title_index]
        meta_lines = entry_buffer[:title_index] + entry_buffer[title_index + 1:]

        title_p = document.add_paragraph()
        title_p.paragraph_format.space_after = Pt(0)
        title_run = title_p.add_run(title_line)
        title_run.bold = True
        title_run.font.size = Pt(11)
        if meta_lines:
            meta_p = document.add_paragraph()
            meta_p.paragraph_format.space_after = Pt(4)
            meta_run = meta_p.add_run(" | ".join(meta_lines))
            meta_run.italic = True
            meta_run.font.size = Pt(10)
            meta_run.font.color.rgb = MUTED
        entry_buffer.clear()

    while i < n:
        stripped = lines[i].strip()

        if _is_section_header(stripped):
            flush_entry()
            heading_p = document.add_paragraph()
            heading_p.paragraph_format.space_before = Pt(14)
            heading_p.paragraph_format.space_after = Pt(4)
            heading_run = heading_p.add_run(stripped.rstrip(":").upper())
            heading_run.bold = True
            heading_run.font.size = Pt(12)
            heading_run.font.color.rgb = ACCENT
            _add_bottom_rule(heading_p, color_hex=str(ACCENT))
            current_section = stripped.rstrip(":").lower()
            i += 1
            continue

        if not stripped:
            flush_entry()
            i += 1
            continue

        if stripped.startswith("- "):
            flush_entry()
            bullet_p = document.add_paragraph(stripped[2:], style="List Bullet")
            bullet_p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        if current_section in ENTRY_SECTIONS:
            entry_buffer.append(stripped)
        else:
            # Summary prose, Skills stray lines, or anything before a section
            # header is recognized: plain paragraph, no bold/italic treatment.
            document.add_paragraph(stripped)
        i += 1

    flush_entry()

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
