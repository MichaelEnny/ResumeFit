"""Parses uploaded resume files (PDF/DOCX) into plain text and flags ATS-risky formatting."""
from dataclasses import dataclass
from io import BytesIO

import pdfplumber
from docx import Document


@dataclass
class ParsedDocument:
    text: str
    warnings: list[str]


def _dedupe(items: list[str]) -> list[str]:
    seen = set()
    result = []
    for item in items:
        if item not in seen:
            seen.add(item)
            result.append(item)
    return result


def parse_pdf(file_bytes: bytes) -> ParsedDocument:
    warnings = []
    lines = []
    with pdfplumber.open(BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            if page.find_tables():
                warnings.append("Detected a table. Most ATS parsers cannot read table content reliably.")
            if page.images:
                warnings.append("Detected an image or graphic. ATS parsers ignore or misread these.")
            text = page.extract_text() or ""
            lines.extend(text.splitlines())
    return ParsedDocument(text="\n".join(lines), warnings=_dedupe(warnings))


def parse_docx(file_bytes: bytes) -> ParsedDocument:
    warnings = []
    document = Document(BytesIO(file_bytes))

    if document.tables:
        warnings.append(f"Detected {len(document.tables)} table(s). Most ATS parsers cannot read table content reliably.")
    if document.inline_shapes:
        warnings.append(f"Detected {len(document.inline_shapes)} image(s) or graphic(s). ATS parsers ignore or misread these.")
    if "txbxContent" in document.element.xml:
        warnings.append("Detected a text box. Most ATS parsers skip text box content entirely.")

    lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    return ParsedDocument(text="\n".join(lines), warnings=_dedupe(warnings))


def parse_uploaded_file(filename: str, file_bytes: bytes) -> ParsedDocument:
    lowered = filename.lower()
    if lowered.endswith(".pdf"):
        return parse_pdf(file_bytes)
    if lowered.endswith(".docx"):
        return parse_docx(file_bytes)
    raise ValueError(f"Unsupported file type: {filename}")
